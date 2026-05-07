import docx
import hashlib
import uuid
import os

def get_full_text(file_path):
    doc = docx.Document(file_path)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    return full_text

def generate_id_by_file_name(file_path: str)-> str:
    base_name =os.path.splitext(os.path.basename(file_path))[0]
    name_hash =hashlib.md5(file_path.encode('utf-8')).hexdigest()[:8]
    return f"{base_name}_{name_hash}"

def parse_doc_to_paragraph_chunks(file_path):
    doc_id = get_file_hash(file_path)
    doc = docx.Document(file_path)

    chunks=[]
    paragraph_counter = 0
    buf_paragraph = ""

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if len(text) !=0:
            buf_paragraph += text + " "
        else:
            if buf_paragraph.strip():
                paragraph_counter +=1

                chunk_data = {
                    "document_id": doc_id,
                    "paragraph_id": f"{doc_id}_{paragraph_counter}",
                    "text": buf_paragraph,
                    "metadata":{
                        "source_file": file_path,
                        "paragraph_index": i
                    }
                }
                chunks.append(chunk_data)
                buf_paragraph = ""

    # сохраняет последний абзац, если тот закончился не пустой строкой а просто текстом
    if buf_paragraph.strip():
                paragraph_counter +=1

                chunk_data = {
                    "document_id": doc_id,
                    "paragraph_id": f"{doc_id}_{paragraph_counter}",
                    "text": buf_paragraph,
                    "metadata":{
                        "source_file": file_path,
                        "paragraph_index": i
                    }
                }
                chunks.append(chunk_data)
    return chunks

def get_metadata(raw_chunks_data: list[dict], file_path: str)-> list[dict]:
    """
    Принимаеи сырые чанки с привязанными номерами параграфов и генерирует для них полные метаданные.
    """
    #doc_id = str(uuid.uuid4())[:4]
    doc_id = generate_id_by_file_name(file_path) #hash
    final_documents = []
    for i, chunk in enumerate(raw_chunks_data):
        text = chunk["text"]
        start_para = chunk["start_paragraph"]

        final_documents.append({
            "chunk_text": text,
            "metadata":{
                "doc_id": doc_id,
                "chunk_id":f"{doc_id}_{i+1}",
                "start_paragraph": start_para,
                "source_file": file_path
            }

        })

    return final_documents

def chunk_by_docx(file_path: str, chunk_size: int = 500, overlap_pct: float = 0.2, separators: list[str] = None) -> list[dict]:
    """
    Читает .docx файл потоково, разбивает текст на базовые фрагменты и склеивает их в чанки.
    Возвращает словарь с текстом чанка и номером параграфа, с которого начался чанк.

    Args:
        file_path: Путь к документу.
        chunk_size (int): Целевой максимальный предел длины для одного чанка в символах. по умолчанию 500
        overlap_pct (float): Доля перекрытия (overlap) между соседними чанками (от 0 до 1). По умолчанию 0.2.
        separators (list[str]): Список разделителей текста по умолчанию ["\n\n", "\n", ". ", " ", ""].

    Returns: 
        list[dict]: Массив сформированных текстовых чанков и номер первого параграфа.
    """
    doc = docx.Document(file_path)
    current_chunk = ""
    chunks = []
    overlap_size = int(chunk_size * overlap_pct)

    current_start_para = None
    if separators is None:
        separators = ["\n\n", "\n", ". ", " ", ""]


    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        last_char = text[-1]
        if last_char in ".!?»\"":
            text += "\n"
        else:
            text += " "
        atoms =_split_to_atoms(text, chunk_size, separators)

        for atom in atoms:
            if current_start_para is None:
                            current_start_para = i+1

            if len(current_chunk) + len(atom) <= chunk_size:
                current_chunk += atom

            else:
                if current_chunk:
                    chunks.append({
                        "text": current_chunk.strip(),
                        "start_paragraph": current_start_para
                                   })
                if overlap_size == 0:
                    overlap_text = ""
                else:
                    raw_ov = current_chunk[-overlap_size:]

                    break_point = -1
                    for s in ["\n", " ", "."]:
                        pos = raw_ov.find(s)
                        if pos != -1 and (break_point == -1 or pos < break_point):
                            break_point = pos
                    
                    overlap_text = raw_ov[break_point:].lstrip() if break_point != -1 else raw_ov
                current_chunk = overlap_text + atom
                current_start_para = i+1

    if current_chunk:
        last_str = current_chunk.strip()
        if chunks and len(last_str) < (chunk_size * 0.25):
            chunks[-1]["text"] = chunks[-1]["text"] + "\n" + last_str
        else:
            chunks.append({
                "text": last_str,
                "start_paragraph": current_start_para
                           })
            
    return chunks

def _split_to_atoms(text: str, chunk_size: int, separators: list[str]) -> list[str]:
    """
    Рекурсивно дробит строку на неделимые элементы.
    """

    if len(text) <= chunk_size or not separators:
        return [text]
    
    sep = separators[0]
    if sep not in text:
        return _split_to_atoms(text, chunk_size, separators[1:])
    
    parts = text.split(sep)
    atoms = []
    for p in parts:
        if len(p) > chunk_size:
            atoms.extend(_split_to_atoms(p, chunk_size, separators[1:]))
        elif p.strip():
            atoms.append(p + sep)
    return atoms

def get_chunks(text: str, chunk_size: int = 500, overlap_pct: float = 0.2) -> list[str]:
    """
    вызывает рекурсивную функцию разбиения текста, а затем склеивает получившиеся фрагменты в чанки.
    Args:
        text (str): Исходная строка текста.
        chunk_size (int): Целевой максимальный предел длины для одного чанка в символах. по умолчанию 500
        overlap_pct (float): Доля перекрытия (overlap) между соседними чанками (от 0 до 1). По умолчанию 0.2.

    Returns: 
        list[str]: Массив сформированных текстовых чанков.
    """

    separators = ["\n\n", "\n", ". ", " ", ""]
    overlap_size = int(chunk_size * overlap_pct)

    atoms = _split_to_atoms(text, chunk_size, separators)
    
    chunks = []
    current_chunk = ""
    

    for atom in atoms:
        if len(current_chunk) + len(atom) <= chunk_size:
            current_chunk += atom
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            
            raw_ov = current_chunk[-overlap_size:]

            break_point = -1
            for s in ["\n", " ", "."]:
                pos = raw_ov.find(s)
                if pos != -1 and (break_point == -1 or pos < break_point):
                    break_point = pos
            
            overlap_text = raw_ov[break_point:].lstrip() if break_point != -1 else raw_ov
            current_chunk = overlap_text + atom

    if current_chunk:
        last_str = current_chunk.strip()
        if chunks and len(last_str) < (chunk_size * 0.25):
            chunks[-1] = chunks[-1] + "\n" + last_str
        else:
            chunks.append(last_str)
            
    return chunks

def process_folder(folder_path: str, chunk_size: int = 500, overlap_pct: float = 0.2, separators: list[str]= None):
    """
    Сканирует папку, обрабатывает все docx файлы и собирает данные 
    для пакетной загрузки в ChromaDB.
    """

    all_texts=[]
    all_metadatas =[]
    all_ids = []
    files = [f for f in os.listdir(folder_path) if f.endswith('.docx')]

    for file_name in files:
        file_path = os.path.join(folder_path, file_name)
        raw_chunks = chunk_by_docx(file_path, chunk_size, overlap_pct, separators)
        process_chunks = get_metadata(raw_chunks, file_path)

        for chunk in process_chunks:
            all_texts.append(chunk['chunk_text'])
            all_metadatas.append(chunk['metadata'])
            all_ids.append(chunk['metadata']['chunk_id'])

    return all_texts, all_metadatas, all_ids

if __name__ == "__main__":
    #path = r'C:\Users\user\Documents\A_dinner_party.docx'
    path = r'Kursach.docx'

    # raw_text = get_full_text(path)

    # chunks = get_chunks(raw_text)
    raw_chunk = chunk_by_docx(path)
    chunks = get_metadata(raw_chunk, path)
    print(f"ОБРАБОТКА ФАЙЛА: {path}")
    print(f"Количество созданных чанков: {len(chunks)}")
    print("=" * 50)

    for i, chunk_obj in enumerate(chunks):
        chunk_text = chunk_obj["chunk_text"]
        metadata = chunk_obj["metadata"]
        print(f"ЧАНК #{i + 1} | Размер: {len(chunk_text)} симв. metadata: {metadata}")
        print("-" * 30)
        print(chunk_text)
        print("=" * 50)